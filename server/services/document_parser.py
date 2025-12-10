"""
Multi-format Document Parser for Data Smith.
Supports PDF, DOC, DOCX, CSV, XLSX, and TXT file formats.
"""

import os
from pathlib import Path
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, field
from enum import Enum
import tempfile
import shutil

# PDF parsing
try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Spreadsheet parsing
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


class FileType(str, Enum):
    """Supported file types."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    CSV = "csv"
    XLSX = "xlsx"
    XLS = "xls"
    UNKNOWN = "unknown"


@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""
    filename: str
    file_type: FileType
    file_size: int
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    char_count: Optional[int] = None
    title: Optional[str] = None
    author: Optional[str] = None
    extra: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentContent:
    """Parsed document content with metadata."""
    text: str
    metadata: DocumentMetadata
    tables: List[Dict[str, Any]] = field(default_factory=list)
    sections: List[Dict[str, str]] = field(default_factory=list)


class DocumentParserError(Exception):
    """Exception raised when document parsing fails."""
    pass


class DocumentParser:
    """
    Multi-format document parser.
    
    Supports:
    - PDF (via pdfplumber or pypdf)
    - DOCX (via python-docx)
    - DOC (limited support via conversion)
    - CSV/XLSX/XLS (via pandas)
    - TXT (native Python)
    """
    
    def __init__(self, temp_dir: Optional[str] = None):
        """
        Initialize the document parser.
        
        Args:
            temp_dir: Directory for temporary files (optional)
        """
        self.temp_dir = temp_dir or tempfile.gettempdir()
        self._validate_dependencies()
    
    def _validate_dependencies(self) -> Dict[str, bool]:
        """Check which parsing libraries are available."""
        return {
            "pdf": PDF_AVAILABLE or PYPDF_AVAILABLE,
            "docx": DOCX_AVAILABLE,
            "spreadsheet": PANDAS_AVAILABLE,
        }
    
    def detect_file_type(self, filename: str) -> FileType:
        """Detect file type from filename extension."""
        ext = Path(filename).suffix.lower().lstrip(".")
        
        type_map = {
            "pdf": FileType.PDF,
            "docx": FileType.DOCX,
            "doc": FileType.DOC,
            "txt": FileType.TXT,
            "csv": FileType.CSV,
            "xlsx": FileType.XLSX,
            "xls": FileType.XLS,
        }
        
        return type_map.get(ext, FileType.UNKNOWN)
    
    async def parse(
        self, 
        file_path: str = None,
        file_content: bytes = None,
        filename: str = None
    ) -> DocumentContent:
        """
        Parse a document and extract its content.
        
        Args:
            file_path: Path to the file on disk
            file_content: Raw file bytes (for uploaded files)
            filename: Original filename (required if using file_content)
            
        Returns:
            DocumentContent with extracted text and metadata
        """
        # Determine source
        if file_content is not None:
            if not filename:
                raise DocumentParserError("filename required when using file_content")
            # Save to temp file for processing
            file_path = self._save_temp_file(file_content, filename)
            cleanup_temp = True
        else:
            if not file_path:
                raise DocumentParserError("Either file_path or file_content required")
            filename = Path(file_path).name
            cleanup_temp = False
        
        try:
            file_type = self.detect_file_type(filename)
            file_size = os.path.getsize(file_path)
            
            # Parse based on file type
            if file_type == FileType.PDF:
                text, tables, extra = await self._parse_pdf(file_path)
            elif file_type in (FileType.DOCX, FileType.DOC):
                text, tables, extra = await self._parse_docx(file_path)
            elif file_type == FileType.TXT:
                text, tables, extra = await self._parse_txt(file_path)
            elif file_type in (FileType.CSV, FileType.XLSX, FileType.XLS):
                text, tables, extra = await self._parse_spreadsheet(file_path, file_type)
            else:
                raise DocumentParserError(f"Unsupported file type: {file_type}")
            
            # Build metadata
            metadata = DocumentMetadata(
                filename=filename,
                file_type=file_type,
                file_size=file_size,
                word_count=len(text.split()) if text else 0,
                char_count=len(text) if text else 0,
                extra=extra
            )
            
            return DocumentContent(
                text=text,
                metadata=metadata,
                tables=tables
            )
            
        finally:
            if cleanup_temp and file_path:
                try:
                    os.unlink(file_path)
                except OSError:
                    pass
    
    def _save_temp_file(self, content: bytes, filename: str) -> str:
        """Save file content to a temporary file."""
        ext = Path(filename).suffix
        temp_path = os.path.join(self.temp_dir, f"datasmith_{os.urandom(8).hex()}{ext}")
        
        with open(temp_path, "wb") as f:
            f.write(content)
        
        return temp_path
    
    async def _parse_pdf(self, file_path: str) -> tuple[str, List[Dict], Dict]:
        """Parse PDF document."""
        text_parts = []
        tables = []
        extra = {}
        
        if PDF_AVAILABLE:
            # Use pdfplumber for better table extraction
            with pdfplumber.open(file_path) as pdf:
                extra["page_count"] = len(pdf.pages)
                
                for i, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table in page_tables:
                        if table:
                            # Convert to dict format
                            headers = table[0] if table else []
                            rows = table[1:] if len(table) > 1 else []
                            tables.append({
                                "page": i + 1,
                                "headers": headers,
                                "rows": rows
                            })
        
        elif PYPDF_AVAILABLE:
            # Fallback to pypdf (text only)
            reader = PdfReader(file_path)
            extra["page_count"] = len(reader.pages)
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        else:
            raise DocumentParserError(
                "PDF parsing not available. Install pdfplumber or pypdf."
            )
        
        return "\n\n".join(text_parts), tables, extra
    
    async def _parse_docx(self, file_path: str) -> tuple[str, List[Dict], Dict]:
        """Parse DOCX document."""
        if not DOCX_AVAILABLE:
            raise DocumentParserError(
                "DOCX parsing not available. Install python-docx."
            )
        
        doc = DocxDocument(file_path)
        text_parts = []
        tables = []
        extra = {}
        
        # Extract core properties
        if doc.core_properties:
            extra["title"] = doc.core_properties.title
            extra["author"] = doc.core_properties.author
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                tables.append({
                    "index": i,
                    "headers": table_data[0] if table_data else [],
                    "rows": table_data[1:] if len(table_data) > 1 else []
                })
        
        return "\n\n".join(text_parts), tables, extra
    
    async def _parse_txt(self, file_path: str) -> tuple[str, List[Dict], Dict]:
        """Parse plain text file."""
        encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                return text, [], {"encoding": encoding}
            except UnicodeDecodeError:
                continue
        
        raise DocumentParserError("Could not decode text file with any supported encoding")
    
    async def _parse_spreadsheet(
        self, 
        file_path: str, 
        file_type: FileType
    ) -> tuple[str, List[Dict], Dict]:
        """Parse CSV/Excel spreadsheet."""
        if not PANDAS_AVAILABLE:
            raise DocumentParserError(
                "Spreadsheet parsing not available. Install pandas and openpyxl."
            )
        
        tables = []
        text_parts = []
        extra = {}
        
        try:
            if file_type == FileType.CSV:
                df = pd.read_csv(file_path)
                sheets = {"Sheet1": df}
            else:
                # Excel file - read all sheets
                excel_file = pd.ExcelFile(file_path)
                sheets = {
                    sheet: pd.read_excel(excel_file, sheet_name=sheet)
                    for sheet in excel_file.sheet_names
                }
                extra["sheet_count"] = len(sheets)
            
            for sheet_name, df in sheets.items():
                # Convert to table format
                tables.append({
                    "sheet": sheet_name,
                    "headers": df.columns.tolist(),
                    "rows": df.values.tolist(),
                    "row_count": len(df)
                })
                
                # Convert to text representation
                text_parts.append(f"### {sheet_name}\n")
                text_parts.append(df.to_string(index=False))
                text_parts.append("\n")
            
        except Exception as e:
            raise DocumentParserError(f"Failed to parse spreadsheet: {e}")
        
        return "\n".join(text_parts), tables, extra


# Convenience function for quick parsing
async def parse_document(
    file_path: str = None,
    file_content: bytes = None,
    filename: str = None
) -> DocumentContent:
    """
    Quick function to parse a document.
    
    Args:
        file_path: Path to file on disk
        file_content: Raw file bytes
        filename: Original filename
        
    Returns:
        DocumentContent with extracted text and metadata
    """
    parser = DocumentParser()
    return await parser.parse(
        file_path=file_path,
        file_content=file_content,
        filename=filename
    )
